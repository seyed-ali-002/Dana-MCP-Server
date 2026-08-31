from __future__ import annotations

from pathlib import Path
from typing import Any

from mcp.server.fastmcp import FastMCP


def _set_rtl(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph.alignment = 2
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def _set_font(run, font: str) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(
        qn("w:cs"), font
    ) if run._element.rPr is not None else None


def register_document_tools(mcp: FastMCP) -> None:
    @mcp.tool()
    def create_docx(
        path: str, title: str = "", content: str = "", font: str = "Tahoma"
    ) -> dict[str, Any]:
        """Create a UTF-8 Persian/RTL-compatible Word document."""
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt

        doc = Document()
        sections = doc.sections
        for section in sections:
            section.right_margin = section.right_margin
        if title:
            p = doc.add_heading(level=1)
            _set_rtl(p)
            r = p.add_run(title)
            r.font.name = font
            r.font.size = Pt(18)
        for block in content.split("\n"):
            p = doc.add_paragraph()
            _set_rtl(p)
            r = p.add_run(block)
            r.font.name = font
            r.font.size = Pt(11)
            rPr = r._element.get_or_add_rPr()
            rFonts = rPr.rFonts or OxmlElement("w:rFonts")
            rFonts.set(qn("w:ascii"), font)
            rFonts.set(qn("w:hAnsi"), font)
            rFonts.set(qn("w:cs"), font)
            if rPr.rFonts is None:
                rPr.append(rFonts)
        out = Path(path).expanduser().resolve()
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(out)
        return {"path": str(out), "format": "docx", "rtl": True, "font": font}

    @mcp.tool()
    def create_pdf(
        path: str, title: str = "", content: str = "", font_path: str | None = None
    ) -> dict[str, Any]:
        """Create a Persian-capable PDF using ReportLab with Arabic shaping and bidi."""
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_RIGHT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        try:
            import arabic_reshaper
            from bidi.algorithm import get_display
        except ImportError as exc:
            raise RuntimeError(
                "Install arabic-reshaper and python-bidi for Persian PDF support"
            ) from exc
        out = Path(path).expanduser().resolve()
        out.parent.mkdir(parents=True, exist_ok=True)
        if not font_path:
            candidates = [
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
            ]
            font_path = next((p for p in candidates if Path(p).exists()), None)
        if not font_path:
            raise RuntimeError("A Unicode TTF font is required; provide font_path")
        font_name = "DanaUnicode"
        pdfmetrics.registerFont(TTFont(font_name, font_path))
        style = ParagraphStyle(
            "DanaRTL",
            fontName=font_name,
            fontSize=11,
            leading=18,
            alignment=TA_RIGHT,
            textColor=colors.black,
        )
        doc = SimpleDocTemplate(
            str(out),
            pagesize=A4,
            rightMargin=50,
            leftMargin=50,
            topMargin=50,
            bottomMargin=50,
        )
        story = []
        for text in ([title] if title else []) + content.split("\n"):
            shaped = get_display(arabic_reshaper.reshape(text))
            story.append(Paragraph(shaped.replace("&", "&amp;"), style))
            story.append(Spacer(1, 8))
        doc.build(story)
        return {"path": str(out), "format": "pdf", "rtl": True, "font": font_name}

    @mcp.tool()
    def create_document(
        path: str,
        content: str,
        title: str = "",
        format: str = "docx",
        font: str = "Tahoma",
        font_path: str | None = None,
    ) -> dict[str, Any]:
        """Create a Word or PDF document with Unicode/RTL support."""
        if format.lower() == "pdf":
            return create_pdf(path, title, content, font_path)
        return create_docx(path, title, content, font)
